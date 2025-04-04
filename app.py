import os
from flask import Flask, render_template, request, redirect, url_for, send_file
from vector_store import create_vector_store, load_vector_store, find_document_type
from PyPDF2 import PdfReader
import logging
import yaml
from groq import Groq
import requests
from docx import Document

# logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)

documents_dir = os.path.join(os.path.dirname(__file__), "documents")
vector_store_path = os.path.join(os.path.dirname(__file__), "vector_store")
# vector_store, total_characters, total_csvs, split_documents, total_documents = create_vector_store(documents_dir, vector_store_path)
vector_store = load_vector_store(vector_store_path)

with open("config.yaml", "r") as file:
    config = yaml.safe_load(file)

with open("prompt.txt", "r") as file:
    prompt = file.read()

groq_api_key = config["GROQ_API_KEY"]
model = "llama-3.3-70b-versatile"
groq_client = Groq(api_key=groq_api_key)

@app.route('/')
def index():
    document_type = request.args.get('document_type')
    clause_names = request.args.getlist('clause_names')
    return render_template('index.html', document_type=document_type, clause_names=clause_names)

@app.route('/upload', methods=['POST'])
def upload():
    if 'contract' not in request.files:
        return redirect(url_for('index'))
    
    file = request.files['contract']
    if file.filename == '':
        return redirect(url_for('index'))
    
    if file and file.filename.endswith('.pdf'):
        file_path = os.path.join(documents_dir, file.filename)
        file.save(file_path)
        pdf_text = extract_text_from_pdf(file_path)
        document_type, clauses = extract_document_info_with_groq(pdf_text, input_text, output_text)
        docx_path = os.path.join(documents_dir, "extract.docx")
        generate_docx_file(docx_path, document_type, clauses)
        return send_file(docx_path, as_attachment=True)
    
    return redirect(url_for('index'))

def generate_docx_file(docx_path, document_type, clauses):
    document = Document()
    document.add_heading('Extracted Information', level=1)
    document.add_heading('Document Type', level=2)
    document.add_paragraph(document_type)
    document.add_heading('Clauses', level=2)

    clauses_list = list(clauses.items())  
    if clauses_list and "Document Type" in clauses_list[0][0]:
        clauses_list.pop(0)

    for clause_name, clause_content in clauses_list:
        document.add_paragraph(f"{clause_name} - {clause_content}")

    document.save(docx_path)

def extract_text_from_pdf(file_path):
    pdf_text = ""
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            pdf_text += page.extract_text()
    return pdf_text

def convert_docx_to_text(docx_path):
    document = Document(docx_path)
    text = "\n".join([para.text for para in document.paragraphs])
    return text

input_file_path = "./finetune/input1.pdf"
output_file_path = "./finetune/output1.docx"
input_text = extract_text_from_pdf(input_file_path)
output_text = convert_docx_to_text(output_file_path)
# logging.debug(f"Input Text: {input_text}")
# logging.debug(f"Output Text: {output_text}")

def extract_document_info_with_groq(pdf_text, input_text, output_text):
    prompt_text = f"""
    The following is an example of a contract clause extraction task. We have an input contract text and an expected output text with extracted clauses. Your goal is to extract the clauses from the provided contract text and match the structure of the expected output text.

    Example:
    Input Contract Text:
    {input_text}

    Expected Output:
    {output_text}

    Your task:
    Given the following contract text:
    {pdf_text}

    Please extract the clauses and their content in the format below:
    - Document Type: <document_type>
    - Clause 1: <Clause Name> - <Clause Content>
    - Clause 2: <Clause Name> - <Clause Content>
    - ...

    If clause names are explicitly mentioned before each clause description, extract those names directly.
    If clause names are not explicitly mentioned, identify the key idea of each clause and match it with relevant clause names from a legal/contractual context. Use the most relevant clause names based on similarity.
    But the clause description should be taken only from the document.
    """

    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt_text}],
            model=model,
            stream=False
        )

        message_content = chat_completion.choices[0].message.content.strip()
        # logging.debug(f"Full Response: {message_content}")

        document_type = ""
        clauses = {}

        if "Document Type:" in message_content:
            start_idx = message_content.find("Document Type:") + len("Document Type:")
            end_idx = message_content.find("\n", start_idx)
            document_type = message_content[start_idx:end_idx].strip()

        clause_lines = message_content.split("- Clause")
        for clause_line in clause_lines:
            if clause_line.strip(): 
                parts = clause_line.split(":", 1)
                if len(parts) == 2:
                    clause_name = parts[0].strip()
                    clause_content = parts[1].strip()
                    clauses[clause_name] = clause_content

        # logging.debug(f"Extracted Document Type: {document_type}")
        # logging.debug(f"Extracted Clauses: {clauses}")

        return document_type, clauses

    except Exception as e:
        # logging.error(f"Error in Groq API call: {str(e)}")
        return 'Error', {}


if __name__ == '__main__':
    app.run(debug=True)